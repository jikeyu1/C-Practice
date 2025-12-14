#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word converter using python-docx and markdown libraries.
Usage: python md_to_docx.py <input_md_file> <output_docx_file>
"""

import sys
import os
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def md_to_docx(md_file_path, docx_file_path):
    """Convert Markdown file to Word document."""
    # Read Markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create a new Word document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "C++面向对象编程实验报告"
    doc.core_properties.author = "Student"
    doc.core_properties.subject = "C++课程实验"
    doc.core_properties.category = "实验报告"
    
    # Convert Markdown to HTML
    html_content = markdown.markdown(md_content)
    
    # Process HTML and add to Word document
    # This is a simplified version - for more complex Markdown, use a better parser
    
    # Split content by lines
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
        
        # Check for headers
        if line.startswith('# '):
            # Level 1 header
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.size = Pt(20)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.bold = True
        elif line.startswith('## '):
            # Level 2 header
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                run.font.size = Pt(16)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.bold = True
        elif line.startswith('### '):
            # Level 3 header
            heading = doc.add_heading(line[4:], level=3)
            for run in heading.runs:
                run.font.size = Pt(14)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.bold = True
        elif line.startswith('#### '):
            # Level 4 header
            heading = doc.add_heading(line[5:], level=4)
            for run in heading.runs:
                run.font.size = Pt(12)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.bold = True
        elif line.startswith('```'):
            # Code block
            if not line[3:]:  # Start of code block
                in_code = True
                code_lines = []
            else:  # End of code block
                in_code = False
                if 'code_lines' in locals() and code_lines:
                    code_paragraph = doc.add_paragraph()
                    code_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in code_paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Courier New'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
                    code_paragraph.add_run('\n'.join(code_lines))
        elif 'in_code' in locals() and locals()['in_code']:
            # Inside code block
            code_lines.append(line)
        elif line.startswith('* '):
            # Bullet list
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.add_run('• ' + line[2:])
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        elif line.startswith('1. '):
            # Numbered list
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.add_run(line)
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        else:
            # Normal text
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Save the Word document
    doc.save(docx_file_path)
    print(f"Successfully converted {md_file_path} to {docx_file_path}")


if __name__ == "__main__":
    # Check if required libraries are installed
    try:
        import markdown
        from docx import Document
    except ImportError as e:
        print(f"Error: Required library not installed - {e}")
        print("Please install the required libraries:")
        print("pip install python-docx markdown")
        sys.exit(1)
    
    # Check command line arguments
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py <input_md_file> <output_docx_file>")
        sys.exit(1)
    
    input_md = sys.argv[1]
    output_docx = sys.argv[2]
    
    # Check if input file exists
    if not os.path.exists(input_md):
        print(f"Error: Input file '{input_md}' not found.")
        sys.exit(1)
    
    # Convert Markdown to Word
    md_to_docx(input_md, output_docx)